#!/usr/bin/env python3
"""
Extract text from DOCX file
"""

import sys
import os

# Try to import docx, install if needed
try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document

def extract_docx_content(filepath):
    """Extract all text from a DOCX file"""
    doc = Document(filepath)
    
    # Extract all paragraphs
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    # Extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

# Extract the document
docx_path = r"C:\Users\scarm\Downloads\Claude Code Integration – Comprehensive Guide.docx"
if os.path.exists(docx_path):
    content = extract_docx_content(docx_path)
    
    # Save as text file
    output_path = r"C:\Users\scarm\llmstack\claude_integration_guide.txt"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"Extracted {len(content)} characters")
    print(f"Saved to: {output_path}")
    print("\n" + "="*60)
    print("CONTENT PREVIEW:")
    print("="*60)
    print(content[:2000])  # Show first 2000 chars
else:
    print(f"File not found: {docx_path}")